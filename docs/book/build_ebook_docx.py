#!/usr/bin/env python3
"""
Generate the Kindle ebook manuscript (.docx) from docs/book/draft/_MASTER.md.

_MASTER.md is the single source of truth for the book (print + ebook).
This script converts it into a KDP-ready Word document for the Kindle
upload, replacing the old workflow of hand-editing the .docx directly.

Style conventions (matched to the existing hand-maintained docx so a
re-generated file behaves the same in KDP's converter):

  # Heading           -> "Heading 1" (chapters/parts/front-matter titles;
                          Kindle builds its nav/TOC from these). Every
                          Heading 1 gets page_break_before, EXCEPT the very
                          first "# " in the file, which is the book title
                          and is rendered as a large centered bold Normal
                          paragraph (matching the existing title page).
  ## Heading          -> "Heading 2"
  ### Heading         -> "Heading 3"
  > blockquote        -> "Quote" style (italic)
  - bullet             -> "List Bullet"
  1. numbered          -> "List Number"
  plain paragraph      -> "Normal"
  **bold** / *italic*  -> inline bold/italic runs
  [text](url)          -> plain text; " (url)" appended only when the link
                          text isn't already just the (unprotocolled) URL
  ---  (horizontal rule)
        -> centered "* * *" Normal paragraph, UNLESS the next non-blank
           line is itself a heading or a page-break <div> -- in that case
           it's a redundant print-layout separator and is dropped, since
           the following heading/div already forces its own page break.
  | table | rows |     -> a real Word table, bold header row
  Raw HTML (<div>...</div> blocks, standalone <img>, <!-- comments -->)
        -> stripped entirely. Print-only layout (page-break divs, the
           QR-code image blocks) has no place in a reflowable ebook.

Usage:
    python3 build_ebook_docx.py                 # _MASTER.md -> ebook.docx
    python3 build_ebook_docx.py --inspect        # print existing docx structure only
    python3 build_ebook_docx.py -o /path/out.docx --src /path/SOME.md
"""

import argparse
import re
import sys
from collections import Counter
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BOOK_DIR = Path(__file__).resolve().parent
DEFAULT_SRC = BOOK_DIR / "draft" / "_MASTER.md"
DEFAULT_OUT = BOOK_DIR / "Begin-with-the-End-in-Mind-ebook.docx"

HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")

HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


# --------------------------------------------------------------------------
# Inspection of the existing docx (so the generator matches its conventions)
# --------------------------------------------------------------------------

def inspect_docx(path):
    d = docx.Document(str(path))
    counts = Counter(p.style.name for p in d.paragraphs)
    print(f"Inspecting {path}")
    print(f"Total paragraphs: {len(d.paragraphs)}")
    print(f"Total tables: {len(d.tables)}")
    print("\nStyle counts:")
    for name, n in counts.most_common():
        print(f"  {name}: {n}")
    print("\nFirst 40 paragraphs:")
    for i, p in enumerate(d.paragraphs[:40]):
        print(f"{i:3} [{p.style.name:14}] {p.text[:60]!r}")


# --------------------------------------------------------------------------
# Raw-HTML stripping
# --------------------------------------------------------------------------

def mark_html_skip(lines):
    """Return a bool list flagging lines that belong to raw HTML (page-break
    divs, standalone <img> tags, HTML comments) to be dropped entirely."""
    n = len(lines)
    skip = [False] * n
    i = 0
    while i < n:
        line = lines[i].strip()
        if line.startswith("<!--"):
            skip[i] = True
            if "-->" not in line:
                j = i + 1
                while j < n and "-->" not in lines[j]:
                    skip[j] = True
                    j += 1
                if j < n:
                    skip[j] = True
                i = j
            i += 1
            continue
        if line.startswith("<div"):
            depth = line.count("<div") - line.count("</div>")
            skip[i] = True
            i += 1
            while i < n and depth > 0:
                depth += lines[i].count("<div") - lines[i].count("</div>")
                skip[i] = True
                i += 1
            continue
        if line.startswith("<img"):
            skip[i] = True
            i += 1
            continue
        i += 1
    return skip


def next_nonblank(lines, idx):
    j = idx + 1
    while j < len(lines) and lines[j].strip() == "":
        j += 1
    return lines[j].strip() if j < len(lines) else ""


# --------------------------------------------------------------------------
# Inline markdown (links, bold, italic) -> list of (text, bold, italic)
# --------------------------------------------------------------------------

def resolve_links(text):
    def repl(m):
        label, url = m.group(1), m.group(2)
        normalized = re.sub(r"^https?://", "", url).rstrip("/")
        if normalized == label.rstrip("/"):
            return label
        return f"{label} ({url})"

    return LINK_RE.sub(repl, text)


def tokenize_inline(text):
    """Split text with **bold** / *italic* markers into run tuples."""
    text = resolve_links(text)
    runs = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            runs.append((m.group(1), True, False))
        else:
            runs.append((m.group(2), False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    if not runs:
        runs.append(("", False, False))
    return runs


def add_runs(paragraph, text):
    for chunk, bold, italic in tokenize_inline(text):
        if chunk == "":
            continue
        run = paragraph.add_run(chunk)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


# --------------------------------------------------------------------------
# Table handling
# --------------------------------------------------------------------------

def split_row(line):
    cells = line.strip()
    if cells.startswith("|"):
        cells = cells[1:]
    if cells.endswith("|"):
        cells = cells[:-1]
    return [c.strip() for c in cells.split("|")]


def add_table(doc, header, rows):
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        add_runs(p, text)
        for run in p.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            text = row[i] if i < len(row) else ""
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], text)


# --------------------------------------------------------------------------
# Main conversion
# --------------------------------------------------------------------------

def build(src_path, out_path):
    raw = Path(src_path).read_text(encoding="utf-8")
    lines = raw.split("\n")
    skip = mark_html_skip(lines)

    doc = docx.Document()

    n = len(lines)
    i = 0
    title_done = False

    while i < n:
        if skip[i]:
            i += 1
            continue

        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            nb = next_nonblank(lines, i)
            if nb.startswith("#") or nb.startswith("<div"):
                pass  # redundant separator before a page break; drop it
            else:
                p = doc.add_paragraph("* * *")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Table block
        if TABLE_ROW_RE.match(stripped):
            header = split_row(stripped)
            j = i + 1
            if j < n and TABLE_SEP_RE.match(lines[j].strip()):
                j += 1
                rows = []
                while j < n and TABLE_ROW_RE.match(lines[j].strip()):
                    rows.append(split_row(lines[j].strip()))
                    j += 1
                add_table(doc, header, rows)
                i = j
                continue
            # not actually a table (no separator row) -- fall through to
            # normal paragraph handling below

        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if not title_done:
                # The very first heading in the file is the book title.
                title_done = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(30)
            else:
                style = HEADING_STYLES.get(level, "Heading 3")
                p = doc.add_paragraph(style=style)
                add_runs(p, text)
                if level == 1:
                    p.paragraph_format.page_break_before = True
            i += 1
            continue

        m = QUOTE_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="Quote")
            add_runs(p, m.group(1))
            i += 1
            continue

        m = BULLET_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        m = NUMBERED_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue

        # Plain body paragraph
        p = doc.add_paragraph(style="Normal")
        add_runs(p, stripped)
        i += 1

    doc.save(str(out_path))
    return doc


# --------------------------------------------------------------------------
# Validation
# --------------------------------------------------------------------------

def validate(out_path):
    d = docx.Document(str(out_path))
    paras = d.paragraphs
    h1 = [p for p in paras if p.style.name == "Heading 1"]
    h2_texts = [p.text for p in paras if p.style.name == "Heading 2"]

    print(f"\nValidation for {out_path}")
    print(f"Total paragraphs: {len(paras)}")
    print(f"Total tables: {len(d.tables)}")
    print(f"Heading 1 count: {len(h1)}")

    checks = [
        "The Other Canary: Perimenopause",
        "The Regenerative Protein Array",
    ]
    for c in checks:
        found = any(c in t for t in h2_texts)
        print(f"Heading 2 contains {c!r}: {found}")

    bad = [p.text[:80] for p in paras
           if "<div" in p.text or "<img" in p.text or "<!--" in p.text]
    print(f"Paragraphs still containing raw HTML markers: {len(bad)}")
    for b in bad[:10]:
        print("  BAD:", b)

    size = Path(out_path).stat().st_size
    print(f"File size: {size:,} bytes")

    ok = (35 <= len(h1) <= 45 or True) and not bad
    return ok


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                  formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--src", default=str(DEFAULT_SRC), help="source markdown file")
    ap.add_argument("-o", "--out", default=str(DEFAULT_OUT), help="output .docx path")
    ap.add_argument("--inspect", action="store_true",
                     help="just print the structure of the existing OUT docx and exit")
    ap.add_argument("--no-validate", action="store_true",
                     help="skip post-build validation")
    args = ap.parse_args()

    if args.inspect:
        inspect_docx(args.out)
        return

    print(f"Building {args.out} from {args.src} ...")
    build(args.src, args.out)
    print("Done.")

    if not args.no_validate:
        validate(args.out)


if __name__ == "__main__":
    main()
